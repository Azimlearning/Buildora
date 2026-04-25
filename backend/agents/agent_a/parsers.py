"""
Agent A: Document Parsers

Handles PDF (digital & scanned) and Word (.docx) documents.

Author: Chip/Azim
"""

import os
import pdfplumber
import fitz  # PyMuPDF
from PIL import Image
import pytesseract
from typing import Optional
import io


def extract_text_pdfplumber(pdf_path: str) -> Optional[str]:
    """
    Extract text from digital PDF using pdfplumber

    Args:
        pdf_path: Path to PDF file

    Returns:
        Extracted text or None if failed
    """
    try:
        with pdfplumber.open(pdf_path) as pdf:
            text = ""
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"

            # Check if we got meaningful text
            if len(text.strip()) > 100:
                return text.strip()

        return None
    except Exception as e:
        print(f"pdfplumber extraction failed: {e}")
        return None


def extract_text_ocr(pdf_path: str) -> str:
    """
    Extract text from scanned PDF using PyMuPDF + Tesseract OCR

    Args:
        pdf_path: Path to PDF file

    Returns:
        Extracted text via OCR
    """
    try:
        doc = fitz.open(pdf_path)
        text = ""

        # Try eng+msa first, fallback to eng if msa is missing
        ocr_lang = "eng+msa"
        try:
            # Quick test to see if msa is available
            pytesseract.get_languages()
            if "msa" not in pytesseract.get_languages():
                ocr_lang = "eng"
        except Exception:
            ocr_lang = "eng"

        for page_num in range(len(doc)):
            page = doc[page_num]

            # Render page to image
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom for better OCR
            img_data = pix.tobytes("png")
            img = Image.open(io.BytesIO(img_data))

            # Run OCR
            try:
                page_text = pytesseract.image_to_string(img, lang=ocr_lang)
            except pytesseract.TesseractError:
                # If eng+msa failed, fallback to eng
                page_text = pytesseract.image_to_string(img, lang="eng")
                
            text += page_text + "\n"

        doc.close()
        return text.strip()

    except Exception as e:
        print(f"OCR extraction failed: {e}")
        return ""


def extract_text_docx(docx_path: str) -> Optional[str]:
    """
    Extract text from a Word (.docx) document using python-docx.

    Args:
        docx_path: Path to the .docx file

    Returns:
        Extracted text or None if failed
    """
    try:
        from docx import Document
        doc = Document(docx_path)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        # Also extract table text
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)
        text = "\n".join(parts)
        if len(text.strip()) > 50:
            print(f"[OK] Extracted text from DOCX ({len(text)} chars)")
            return text.strip()
        return None
    except ImportError:
        print("[WARN] python-docx not installed. Install with: pip install python-docx")
        return None
    except Exception as e:
        print(f"DOCX extraction failed: {e}")
        return None


def parse_pdf(pdf_path: str) -> str:
    """
    Main PDF parsing function with fallback strategy.
    Alias for parse_document() — kept for backward compatibility.
    """
    return parse_document(pdf_path)


def parse_document(file_path: str) -> str:
    """
    Parse any supported document and return its text content.

    Supported formats:
    - PDF  (.pdf)  — pdfplumber → OCR fallback
    - Word (.docx) — python-docx

    Args:
        file_path: Path to the document

    Returns:
        Extracted text content

    Raises:
        ValueError: If text extraction fails for all strategies
    """
    ext = os.path.splitext(file_path)[1].lower()

    # ── Word document ──────────────────────────────────────────────
    if ext in (".docx", ".doc"):
        text = extract_text_docx(file_path)
        if text:
            return text
        raise ValueError(f"Failed to extract text from Word document: {file_path}")

    # ── PDF (default) ──────────────────────────────────────────────
    # Try digital extraction first
    text = extract_text_pdfplumber(file_path)
    if text:
        print(f"[OK] Extracted text using pdfplumber ({len(text)} chars)")
        return text

    # Fall back to OCR
    print("[WARN] pdfplumber failed, falling back to OCR...")
    text = extract_text_ocr(file_path)
    if text:
        print(f"[OK] Extracted text using OCR ({len(text)} chars)")
        return text

    raise ValueError("Failed to extract text from document using all available methods")


def extract_tables(pdf_path: str) -> list[dict]:
    """
    Extract tables from PDF using pdfplumber

    Args:
        pdf_path: Path to PDF file

    Returns:
        List of table dictionaries
    """
    tables = []

    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                page_tables = page.extract_tables()
                for table in page_tables:
                    tables.append({
                        "page": page_num + 1,
                        "data": table
                    })

    except Exception as e:
        print(f"Table extraction failed: {e}")

    return tables
